from dataclasses import dataclass
from pathlib import Path
import subprocess
import tempfile


class DocumentExtractionError(Exception):
    pass


@dataclass(frozen=True)
class DocumentExtractionResult:
    text: str
    metadata: dict[str, object]


def extract_document_text(path: str, original_filename: str) -> DocumentExtractionResult:
    source_path = Path(path)
    extension = Path(original_filename).suffix.lower() or source_path.suffix.lower()

    if extension == ".pdf":
        return extract_pdf_text(source_path)
    if extension == ".docx":
        return extract_docx_text(source_path)
    if extension == ".doc":
        return extract_doc_text(source_path)
    if extension in {".md", ".txt", ".csv"}:
        return extract_plain_text(source_path, extension)

    raise DocumentExtractionError("Unsupported document type for text extraction")


def extract_pdf_text(path: Path) -> DocumentExtractionResult:
    import fitz

    try:
        with fitz.open(path) as document:
            page_text = [page.get_text("text") for page in document]
            text = normalize_text("\n".join(page_text))
            metadata = {"page_count": document.page_count, "extractor": "pymupdf"}
    except Exception as exc:
        raise DocumentExtractionError("Could not extract selectable PDF text") from exc

    if not text:
        raise DocumentExtractionError("No selectable text found; OCR not supported")
    return DocumentExtractionResult(text=text, metadata=metadata)


def extract_docx_text(path: Path) -> DocumentExtractionResult:
    from docx import Document

    try:
        document = Document(path)
    except Exception as exc:
        raise DocumentExtractionError("Could not read Word document") from exc

    parts: list[str] = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))

    text = normalize_text("\n".join(parts))
    if not text:
        raise DocumentExtractionError("No selectable text found in Word document")
    return DocumentExtractionResult(
        text=text,
        metadata={
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "extractor": "python-docx",
        },
    )


def extract_doc_text(path: Path) -> DocumentExtractionResult:
    with tempfile.TemporaryDirectory() as temp_dir:
        command = [
            "libreoffice",
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            temp_dir,
            str(path),
        ]
        try:
            subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        except FileNotFoundError as exc:
            raise DocumentExtractionError("LibreOffice is required to extract .doc files") from exc
        except subprocess.CalledProcessError as exc:
            detail = exc.stderr.strip() or exc.stdout.strip() or "LibreOffice conversion failed"
            raise DocumentExtractionError(detail[:500]) from exc

        converted = Path(temp_dir) / f"{path.stem}.docx"
        if not converted.exists():
            candidates = list(Path(temp_dir).glob("*.docx"))
            converted = candidates[0] if candidates else converted
        if not converted.exists():
            raise DocumentExtractionError("LibreOffice did not create a .docx file")

        result = extract_docx_text(converted)
        return DocumentExtractionResult(
            text=result.text,
            metadata={**result.metadata, "converted_from": ".doc", "converter": "libreoffice"},
        )


def extract_plain_text(path: Path, extension: str) -> DocumentExtractionResult:
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise DocumentExtractionError("Could not decode text document")

    text = normalize_text(text)
    if not text:
        raise DocumentExtractionError("No text found in document")
    return DocumentExtractionResult(text=text, metadata={"encoding": encoding, "extension": extension})


def normalize_text(text: str) -> str:
    lines = [" ".join(line.split()) for line in text.replace("\x00", "").splitlines()]
    return "\n".join(line for line in lines if line).strip()
