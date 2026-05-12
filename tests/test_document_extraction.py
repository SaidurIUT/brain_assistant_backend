from pathlib import Path

import pytest

from app.services.document_extraction import DocumentExtractionError, extract_document_text


def test_extract_markdown_text(tmp_path: Path) -> None:
    path = tmp_path / "faq.md"
    path.write_text("# FAQ\n\nBrain Assistant supports uploads.", encoding="utf-8")

    result = extract_document_text(str(path), "faq.md")

    assert "Brain Assistant supports uploads." in result.text
    assert result.metadata["extension"] == ".md"


def test_extract_txt_with_latin1_fallback(tmp_path: Path) -> None:
    path = tmp_path / "policy.txt"
    path.write_bytes("Caf\xe9 policy".encode("latin-1"))

    result = extract_document_text(str(path), "policy.txt")

    assert "Café policy" in result.text


def test_extract_docx_text_and_table(tmp_path: Path) -> None:
    from docx import Document

    path = tmp_path / "guide.docx"
    document = Document()
    document.add_paragraph("Welcome to the guide.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Plan"
    table.cell(0, 1).text = "Enterprise"
    document.save(path)

    result = extract_document_text(str(path), "guide.docx")

    assert "Welcome to the guide." in result.text
    assert "Plan | Enterprise" in result.text
    assert result.metadata["extractor"] == "python-docx"


def test_extract_selectable_pdf_text(tmp_path: Path) -> None:
    import fitz

    path = tmp_path / "selectable.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Selectable PDF text")
    document.save(path)
    document.close()

    result = extract_document_text(str(path), "selectable.pdf")

    assert "Selectable PDF text" in result.text
    assert result.metadata["extractor"] == "pymupdf"


def test_empty_pdf_reports_no_selectable_text(tmp_path: Path) -> None:
    import fitz

    path = tmp_path / "empty.pdf"
    document = fitz.open()
    document.new_page()
    document.save(path)
    document.close()

    with pytest.raises(DocumentExtractionError, match="No selectable text found"):
        extract_document_text(str(path), "empty.pdf")
